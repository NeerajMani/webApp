import asyncio
import tkinter as tk
from tkinter import filedialog
from tkinter import messagebox
from datetime import datetime
import docx
import pytz
from pyppeteer import launch

# Function to convert date and time to CST epoch time
def convert_to_epoch(date_string, time_string, timezone):
    cst = pytz.timezone(timezone)
    date_time = datetime.strptime(f"{date_string} {time_string}", "%m/%d/%Y %H:%M")
    date_time_cst = cst.localize(date_time).astimezone(pytz.utc)
    epoch_time = int(date_time_cst.timestamp())
    return epoch_time

# Function to take a screenshot of the graph
async def take_screenshot(page, api_name):
    graph_element = await page.waitForSelector("#splunk-main-content")
    graph_box = await graph_element.boundingBox()
    screenshot_path = f"{api_name}.png"
    await page.screenshot({'path': screenshot_path, 'clip': graph_box})
    return screenshot_path

# Function to update the Word document with the graph
def update_word_document(api_name, graph_image, word_doc):
    word_doc.add_paragraph(api_name)
    word_doc.add_picture(graph_image)

# Function to handle the "Select Document" button click event
def select_document():
    word_doc_path = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
    word_doc_entry.delete(0, tk.END)
    word_doc_entry.insert(0, word_doc_path)

# Function to handle the "Update" button click event
async def update_graphs():
    api_names = api_names_entry.get().split(",")
    domain_name = domain_name_entry.get()
    start_datetime_str = start_datetime_entry.get()
    end_datetime_str = end_datetime_entry.get()
    word_doc_path = word_doc_entry.get()

    # Split start and end date/time strings
    start_date_str, start_time_str = start_datetime_str.split(" ")
    end_date_str, end_time_str = end_datetime_str.split(" ")

    # Convert start and end dates to epoch time
    start_epoch = convert_to_epoch(start_date_str, start_time_str, "US/Central")
    end_epoch = convert_to_epoch(end_date_str, end_time_str, "US/Central")

    # Launch Puppeteer
    async def run():
        browser = await launch(
            executablePath=r'C:\Program Files\Google\Chrome\Application\chrome.exe',
            args=['--no-sandbox', '--start-maximized'],  # Open browser in full screen
            headless=False
        )
        page = await browser.newPage()
        await page.setViewport({'width': 1920, 'height': 1080})  # Set viewport size to full screen

        # Open Splunk in the same tab
        await page.goto("https://www.splunk.com", {"waitUntil": "networkidle0"})

        # Run queries and update Word document
        word_doc = docx.Document(word_doc_path)
        for api_name in api_names:
            # Build the URL with API name, domain, start date, and end date
            url = f"https://www.splunk.com/{api_name}/{domain_name}?start={start_epoch}&end={end_epoch}"
            await page.goto(url, {"waitUntil": "networkidle0"})
            # Wait for the desired object to appear on the screen
            await page.waitForSelector("#object-id", {"timeout": 10000})

            # Take a screenshot of the graph
            graph_image_path_cpu = await take_screenshot(page, api_name)

            # Update the Word document with the API name and graph
            update_word_document(api_name, graph_image_path_cpu, word_doc)

            url = f"https://www.splunk.com/{api_name}/{domain_name}?start={start_epoch}&end={end_epoch}"
            await page.goto(url, {"waitUntil": "networkidle0"})
            # Wait for the desired object to appear on the screen
            await page.waitForSelector("#object-id", {"timeout": 10000})

            # Take a screenshot of the graph
            graph_image_path_mem = await take_screenshot(page, api_name)

            # Update the Word document with the API name and graph
            update_word_document(api_name, graph_image_path_mem, word_doc)

        # Save and close the Word document
        word_doc.save(word_doc_path)

        # Close the browser
        await browser.close()

    # Run the asyncio event loop
    await run()

    messagebox.showinfo("Update Completed", "Graphs updated successfully.")

# Create the main window
window = tk.Tk()
window.title("Graph Update Utility")

# Create UI elements
api_names_label = tk.Label(window, text="API Names (comma-separated):")
api_names_entry = tk.Entry(window, width=50)
domain_name_label = tk.Label(window, text="Domain Name:")
domain_name_entry = tk.Entry(window, width=50)
start_datetime_label = tk.Label(window, text="Start Date/Time (MM/DD/YYYY HH:MM):")
start_datetime_entry = tk.Entry(window, width=50)
end_datetime_label = tk.Label(window, text="End Date/Time (MM/DD/YYYY HH:MM):")
end_datetime_entry = tk.Entry(window, width=50)
word_doc_label = tk.Label(window, text="Word Document:")
word_doc_entry = tk.Entry(window, width=50)
select_doc_button = tk.Button(window, text="Select Document", command=select_document)
update_button = tk.Button(window, text="Update Graphs", command=lambda: asyncio.run(update_graphs()))

# Arrange UI elements using grid layout
api_names_label.grid(row=0, column=0, sticky="W")
api_names_entry.grid(row=0, column=1, columnspan=2)
domain_name_label.grid(row=1, column=0, sticky="W")
domain_name_entry.grid(row=1, column=1, columnspan=2)
start_datetime_label.grid(row=2, column=0, sticky="W")
start_datetime_entry.grid(row=2, column=1, columnspan=2)
end_datetime_label.grid(row=3, column=0, sticky="W")
end_datetime_entry.grid(row=3, column=1, columnspan=2)
word_doc_label.grid(row=4, column=0, sticky="W")
word_doc_entry.grid(row=4, column=1)
select_doc_button.grid(row=4, column=2)
update_button.grid(row=5, column=1, pady=10)

# Start the main loop
window.mainloop()


--------------------------
# Function to take a screenshot of the graph
async def take_screenshot(page, api_name):
    graph_element = await page.waitForSelector("#splunk-main-content")
    graph_box = await graph_element.boundingBox()

    # Calculate the dimensions of the graph element
    graph_width = graph_box['width']
    graph_height = graph_box['height']
    
    # Set the viewport size based on the graph dimensions
    viewport_width = graph_width + 50  # Add some extra width to account for margins, borders, etc.
    viewport_height = graph_height + 100  # Add some extra height to account for margins, borders, etc.
    await page.setViewport({
        'width': viewport_width,
        'height': viewport_height
    })

    # Take the screenshot
    screenshot_path = f"{api_name}.png"
    await page.screenshot({'path': screenshot_path, 'clip': graph_box})
    return screenshot_path

# ...

# Launch Puppeteer and create a new page
browser = await launch_browser()
page = await browser.newPage()

# ...

# Run queries and update Word document
word_doc = docx.Document()
for api_name in api_names:
    # ...

    # Take a screenshot of the graph
    graph_image_path = await take_screenshot(page, api_name)

    # ...

# Save and close the Word document
word_doc.save("output.docx")

-------------------------------------------------
from docx.shared import Inches

# Function to update the Word document with the graph
def update_word_document(api_name, graph_image, word_doc):
    word_doc.add_paragraph(api_name)
    
    # Load the graph image and resize it to fit within the A4 page size
    graph = word_doc.add_picture(graph_image)
    graph.width = Inches(7.5)  # Adjust the width as needed
    graph.height = Inches(5.5)  # Adjust the height as needed

    # Adjust the layout options to position the image within the page
    graph.left = Inches(0.75)
    graph.top = Inches(0.75)
    graph.right = word_doc.page_width - Inches(0.75)
    graph.bottom = word_doc.page_height - Inches(0.75)
    
  -------------------------------

from docx.shared import Inches, Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ...

# Function to update the Word document with the graph
def update_word_document(api_name, graph_image, word_doc):
    # Add a new section to the Word document
    section = word_doc.add_section()

    # Set the page size to A4 format (210mm x 297mm)
    section._sectPr.xpath('./w:pgSz')[0].attrib.clear()
    section._sectPr.xpath('./w:pgSz')[0].attrib.update({
        nsdecls('w'): 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'w:w': '11906',
        'w:h': '16838'
    })

    # Add the API name as a heading
    word_doc.add_heading(api_name, level=1)

    # Add the graph image
    word_doc.add_picture(graph_image, width=Inches(6), height=Inches(4))

# ...

# Run queries and update Word document
word_doc = docx.Document()
for api_name in api_names:
    # ...

    # Take a screenshot of the graph
    graph_image_path = await take_screenshot(page, api_name)

    # Update the Word document with the API name and graph
    update_word_document(api_name, graph_image_path, word_doc)

# ...

# Save and close the Word document
word_doc.save("output.docx")

---------------------------------
# Function to update the Word document with the graph
def update_word_document(api_name, additional_text, graph_image, word_doc):
    # Add a new section to the Word document
    section = word_doc.add_section()

    # Set the page size to A4 format (210mm x 297mm)
    section._sectPr.xpath('./w:pgSz')[0].attrib.clear()
    section._sectPr.xpath('./w:pgSz')[0].attrib.update({
        nsdecls('w'): 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'w:w': '11906',
        'w:h': '16838'
    })

    # Create the heading text by concatenating the API name and additional text
    heading_text = f"{api_name} {additional_text}"

    # Add the heading to the Word document
    word_doc.add_heading(heading_text, level=1)

    # Add the graph image
    word_doc.add_picture(graph_image, width=Inches(6), height=Inches(4))

# ...

# Run queries and update Word document
word_doc = docx.Document()
for api_name in api_names:
    # ...

    # Take a screenshot of the graph
    graph_image_path = await take_screenshot(page, api_name)

    # Additional text for the heading
    additional_text = "Max CPU"

    # Update the Word document with the API name, additional text, and graph
    update_word_document(api_name, additional_text, graph_image_path, word_doc)

# ...

